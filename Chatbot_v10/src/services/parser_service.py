"""
문서 파싱 서비스 - Boundary Layer
계획서 2.1: 구조 보존형 Parser
"""
import re
from pathlib import Path
from typing import Optional, Dict, Any, List
from datetime import datetime

from src.domain.fragments import StructuredDoc, Block
from src.shared.errors import ParsingError
from src.shared.logging import get_logger, log_error
from src.shared.types import Result

logger = get_logger(__name__)


class ParserService:
    """
    Boundary Layer - 구조 보존형 Parser
    
    지원 형식:
    - PDF → Markdown/HTML 변환
    - DOCX → 구조화된 블록
    - HTML/Markdown → 구조 파싱
    - 텍스트 → 단락 분리
    
    원칙 4: 단일 책임 - 문서 파싱만 담당
    """
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".html", ".htm", ".md", ".txt"}
    
    def __init__(self):
        self._parsers = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".doc": self._parse_docx,
            ".html": self._parse_html,
            ".htm": self._parse_html,
            ".md": self._parse_markdown,
            ".txt": self._parse_text,
        }
    
    async def parse_file(self, file_path: str) -> Result[StructuredDoc]:
        """
        파일 파싱
        
        Args:
            file_path: 파일 경로
        
        Returns:
            Result[StructuredDoc]: 구조화된 문서
        """
        path = Path(file_path)
        
        if not path.exists():
            return Result.fail(
                error=f"파일을 찾을 수 없습니다: {file_path}",
                error_code="FILE_NOT_FOUND"
            )
        
        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            return Result.fail(
                error=f"지원하지 않는 파일 형식: {ext}",
                error_code="UNSUPPORTED_FORMAT",
                details={"supported": list(self.SUPPORTED_EXTENSIONS)}
            )
        
        try:
            parser = self._parsers.get(ext)
            if parser:
                doc = await parser(path)
                doc.processed = True
                logger.info(
                    "document_parsed",
                    file_path=file_path,
                    block_count=len(doc.blocks)
                )
                return Result.ok(doc)
            else:
                return Result.fail(
                    error=f"파서를 찾을 수 없습니다: {ext}",
                    error_code="PARSER_NOT_FOUND"
                )
                
        except Exception as e:
            log_error(logger, e, {"file_path": file_path})
            return Result.fail(
                error=str(e),
                error_code="PARSING_FAILED"
            )
    
    async def parse_text_content(self, content: str, source_name: str = "direct_input") -> Result[StructuredDoc]:
        """
        텍스트 콘텐츠 직접 파싱
        
        Args:
            content: 텍스트 내용
            source_name: 출처 이름
        
        Returns:
            Result[StructuredDoc]
        """
        try:
            blocks = self._split_into_blocks(content)
            
            doc = StructuredDoc(
                source_path=source_name,
                source_type="text",
                blocks=blocks,
                metadata={
                    "parsed_at": datetime.utcnow().isoformat(),
                    "char_count": len(content),
                },
                processed=True
            )
            
            return Result.ok(doc)
            
        except Exception as e:
            log_error(logger, e, {"source_name": source_name})
            return Result.fail(error=str(e), error_code="PARSING_FAILED")
    
    async def _parse_pdf(self, path: Path) -> StructuredDoc:
        """PDF 파싱"""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ParsingError("PyPDF2가 설치되지 않았습니다", document_id=str(path))
        
        reader = PdfReader(path)
        blocks: List[Block] = []
        
        metadata = {
            "page_count": len(reader.pages),
            "parsed_at": datetime.utcnow().isoformat(),
        }
        
        # PDF 메타데이터 추출
        if reader.metadata:
            if reader.metadata.title:
                metadata["title"] = reader.metadata.title
            if reader.metadata.author:
                metadata["author"] = reader.metadata.author
            if reader.metadata.creation_date:
                metadata["creation_date"] = str(reader.metadata.creation_date)
        
        order = 0
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            
            # 페이지별 블록 생성
            page_blocks = self._split_into_blocks(text, start_order=order)
            for block in page_blocks:
                block.metadata["page"] = page_num + 1
            
            blocks.extend(page_blocks)
            order += len(page_blocks)
        
        return StructuredDoc(
            source_path=str(path),
            source_type="pdf",
            title=metadata.get("title"),
            blocks=blocks,
            metadata=metadata
        )
    
    async def _parse_docx(self, path: Path) -> StructuredDoc:
        """DOCX 파싱"""
        try:
            from docx import Document
        except ImportError:
            raise ParsingError("python-docx가 설치되지 않았습니다", document_id=str(path))
        
        doc = Document(path)
        blocks: List[Block] = []
        
        metadata = {
            "parsed_at": datetime.utcnow().isoformat(),
        }
        
        # 문서 속성 추출
        if doc.core_properties:
            props = doc.core_properties
            if props.title:
                metadata["title"] = props.title
            if props.author:
                metadata["author"] = props.author
            if props.created:
                metadata["created"] = str(props.created)
        
        order = 0
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            # 스타일로 블록 유형 결정
            block_type = "paragraph"
            level = 0
            
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if "heading" in style_name:
                    block_type = "heading"
                    # Heading 1, Heading 2 등에서 숫자 추출
                    match = re.search(r'\d+', style_name)
                    level = int(match.group()) if match else 1
                elif "list" in style_name:
                    block_type = "list"
            
            blocks.append(Block(
                block_type=block_type,
                content=para.text.strip(),
                order=order,
                level=level
            ))
            order += 1
        
        # 테이블 파싱
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_content.append(" | ".join(row_data))
            
            if table_content:
                blocks.append(Block(
                    block_type="table",
                    content="\n".join(table_content),
                    order=order,
                    metadata={"row_count": len(table.rows)}
                ))
                order += 1
        
        return StructuredDoc(
            source_path=str(path),
            source_type="docx",
            title=metadata.get("title"),
            blocks=blocks,
            metadata=metadata
        )
    
    async def _parse_html(self, path: Path) -> StructuredDoc:
        """HTML 파싱"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ParsingError("beautifulsoup4가 설치되지 않았습니다", document_id=str(path))
        
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        
        soup = BeautifulSoup(content, "html.parser")
        blocks: List[Block] = []
        
        # 제목 추출
        title = None
        title_tag = soup.find("title")
        if title_tag:
            title = title_tag.get_text().strip()
        
        metadata = {
            "parsed_at": datetime.utcnow().isoformat(),
            "title": title,
        }
        
        order = 0
        
        # 헤딩 태그
        for level in range(1, 7):
            for heading in soup.find_all(f"h{level}"):
                text = heading.get_text().strip()
                if text:
                    blocks.append(Block(
                        block_type="heading",
                        content=text,
                        order=order,
                        level=level
                    ))
                    order += 1
        
        # 단락
        for p in soup.find_all("p"):
            text = p.get_text().strip()
            if text:
                blocks.append(Block(
                    block_type="paragraph",
                    content=text,
                    order=order
                ))
                order += 1
        
        # 리스트
        for ul in soup.find_all(["ul", "ol"]):
            items = [li.get_text().strip() for li in ul.find_all("li")]
            if items:
                blocks.append(Block(
                    block_type="list",
                    content="\n".join(f"- {item}" for item in items),
                    order=order
                ))
                order += 1
        
        # 테이블
        for table in soup.find_all("table"):
            rows = []
            for tr in table.find_all("tr"):
                cells = [td.get_text().strip() for td in tr.find_all(["td", "th"])]
                rows.append(" | ".join(cells))
            
            if rows:
                blocks.append(Block(
                    block_type="table",
                    content="\n".join(rows),
                    order=order
                ))
                order += 1
        
        # order로 정렬
        blocks.sort(key=lambda b: b.order)
        
        return StructuredDoc(
            source_path=str(path),
            source_type="html",
            title=title,
            blocks=blocks,
            metadata=metadata
        )
    
    async def _parse_markdown(self, path: Path) -> StructuredDoc:
        """Markdown 파싱"""
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        
        blocks: List[Block] = []
        lines = content.split("\n")
        
        metadata = {
            "parsed_at": datetime.utcnow().isoformat(),
        }
        
        order = 0
        current_block = []
        current_type = "paragraph"
        current_level = 0
        
        for line in lines:
            stripped = line.strip()
            
            # 헤딩
            if stripped.startswith("#"):
                # 이전 블록 저장
                if current_block:
                    blocks.append(Block(
                        block_type=current_type,
                        content="\n".join(current_block).strip(),
                        order=order,
                        level=current_level
                    ))
                    order += 1
                    current_block = []
                
                # 헤딩 레벨 계산
                level = len(stripped) - len(stripped.lstrip("#"))
                text = stripped.lstrip("#").strip()
                
                blocks.append(Block(
                    block_type="heading",
                    content=text,
                    order=order,
                    level=level
                ))
                order += 1
                current_type = "paragraph"
                current_level = 0
            
            # 리스트
            elif stripped.startswith(("-", "*", "+")) or re.match(r"^\d+\.", stripped):
                if current_type != "list" and current_block:
                    blocks.append(Block(
                        block_type=current_type,
                        content="\n".join(current_block).strip(),
                        order=order,
                        level=current_level
                    ))
                    order += 1
                    current_block = []
                
                current_type = "list"
                current_block.append(stripped)
            
            # 빈 줄
            elif not stripped:
                if current_block:
                    blocks.append(Block(
                        block_type=current_type,
                        content="\n".join(current_block).strip(),
                        order=order,
                        level=current_level
                    ))
                    order += 1
                    current_block = []
                    current_type = "paragraph"
                    current_level = 0
            
            # 일반 텍스트
            else:
                if current_type == "list":
                    blocks.append(Block(
                        block_type=current_type,
                        content="\n".join(current_block).strip(),
                        order=order
                    ))
                    order += 1
                    current_block = []
                    current_type = "paragraph"
                
                current_block.append(stripped)
        
        # 마지막 블록 저장
        if current_block:
            blocks.append(Block(
                block_type=current_type,
                content="\n".join(current_block).strip(),
                order=order,
                level=current_level
            ))
        
        # 제목 추출 (첫 번째 H1)
        title = None
        for block in blocks:
            if block.block_type == "heading" and block.level == 1:
                title = block.content
                break
        
        return StructuredDoc(
            source_path=str(path),
            source_type="markdown",
            title=title,
            blocks=blocks,
            metadata=metadata
        )
    
    async def _parse_text(self, path: Path) -> StructuredDoc:
        """텍스트 파일 파싱"""
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        
        blocks = self._split_into_blocks(content)
        
        return StructuredDoc(
            source_path=str(path),
            source_type="text",
            blocks=blocks,
            metadata={
                "parsed_at": datetime.utcnow().isoformat(),
                "char_count": len(content),
            }
        )
    
    def _split_into_blocks(self, text: str, start_order: int = 0) -> List[Block]:
        """텍스트를 단락 블록으로 분리"""
        blocks: List[Block] = []
        
        # 빈 줄로 단락 분리
        paragraphs = re.split(r"\n\s*\n", text)
        
        order = start_order
        for para in paragraphs:
            para = para.strip()
            if para:
                blocks.append(Block(
                    block_type="paragraph",
                    content=para,
                    order=order
                ))
                order += 1
        
        return blocks
